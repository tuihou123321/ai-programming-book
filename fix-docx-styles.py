#!/usr/bin/env python3
"""
修复DOCX文件样式以符合中文排版要求
需要安装: pip install python-docx
"""

import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def fix_docx_styles(input_file, output_file):
    """修复DOCX文件的样式"""
    try:
        # 打开文档
        doc = Document(input_file)
        
        # 设置字体映射
        def set_font(run, font_name, size_pt, bold=False):
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 设置中文字体
            run.font.size = Pt(size_pt)
            run.font.bold = bold
        
        # 遍历所有段落
        for paragraph in doc.paragraphs:
            # 根据样式设置字体
            if paragraph.style.name.startswith('Heading 1'):
                # 一级标题：小二(18pt)、宋体、加粗
                for run in paragraph.runs:
                    set_font(run, '宋体', 18, True)
                    
            elif paragraph.style.name.startswith('Heading 2'):
                # 二级标题：小三(15pt)、黑体（不加粗）
                for run in paragraph.runs:
                    set_font(run, '黑体', 15, False)
                    
            elif paragraph.style.name.startswith('Heading 3'):
                # 三级标题：小四(12pt)、黑体（不加粗）
                for run in paragraph.runs:
                    set_font(run, '黑体', 12, False)
                    
            elif paragraph.style.name == 'Source Code':
                # 代码块：小五(9pt)、Courier New
                for run in paragraph.runs:
                    set_font(run, 'Courier New', 9, False)
                    
            elif paragraph.style.name == 'Block Text':
                # 引用块（提示）：黑体、五号(9pt)
                for run in paragraph.runs:
                    set_font(run, '黑体', 9, False)
                    
            else:
                # 处理正文和行内代码
                for run in paragraph.runs:
                    # 检查run的样式，寻找行内代码
                    run_style = None
                    try:
                        # 尝试获取run的样式
                        if hasattr(run, '_element') and run._element.rPr is not None:
                            style_elem = run._element.rPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rStyle')
                            if style_elem is not None:
                                run_style = style_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    except:
                        pass
                    
                    # 如果是行内代码样式
                    if run_style == 'VerbatimChar':
                        # 行内代码：小五(9pt)、Courier New
                        set_font(run, 'Courier New', 9, False)
                    elif not run.font.name or run.font.name == 'Calibri':
                        # 正文：宋体、10号
                        set_font(run, '宋体', 10, False)
        
        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_font(run, '宋体', 10, False)
        
        # 保存修改后的文档
        doc.save(output_file)
        print(f"样式修复完成: {output_file}")
        
    except Exception as e:
        print(f"修复样式时出错: {e}")
        return False
    
    return True

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python fix-docx-styles.py <输入文件> <输出文件>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    success = fix_docx_styles(input_file, output_file)
    if success:
        print("✅ DOCX样式修复成功!")
    else:
        print("❌ DOCX样式修复失败!")
        sys.exit(1)