#!/usr/bin/env python3
"""
调试DOCX文件样式，查看Pandoc生成的实际样式名称
"""

import sys
from docx import Document

def debug_docx_styles(input_file):
    """调试DOCX文件的样式"""
    try:
        doc = Document(input_file)
        
        print("=== DOCX样式调试信息 ===\n")
        
        # 打印所有可用样式
        print("📋 文档中的所有样式:")
        for style in doc.styles:
            print(f"  - {style.name} ({style.type})")
        
        print("\n" + "="*50 + "\n")
        
        # 遍历段落并显示样式
        print("📝 段落样式分析:")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 只显示非空段落
                print(f"\n段落 {i+1}: '{paragraph.text[:50]}{'...' if len(paragraph.text) > 50 else ''}'")
                print(f"  样式名: {paragraph.style.name}")
                
                # 显示每个run的字体信息
                for j, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        font_name = run.font.name or "未设置"
                        font_size = run.font.size.pt if run.font.size else "未设置"
                        is_bold = run.font.bold
                        print(f"    Run {j+1}: '{run.text[:30]}{'...' if len(run.text) > 30 else ''}'")
                        print(f"      字体: {font_name}, 大小: {font_size}pt, 加粗: {is_bold}")
        
        print("\n" + "="*50 + "\n")
        
        # 分析表格
        if doc.tables:
            print("📊 表格样式分析:")
            for i, table in enumerate(doc.tables):
                print(f"\n表格 {i+1}:")
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if paragraph.text.strip():
                                print(f"  行{row_idx+1}列{cell_idx+1}: {paragraph.style.name}")
                                break
                        break
                    if row_idx >= 2:  # 只显示前3行
                        break
        
    except Exception as e:
        print(f"调试时出错: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("用法: python debug-docx-styles.py <DOCX文件>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    debug_docx_styles(input_file)